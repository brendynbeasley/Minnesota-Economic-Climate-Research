"""Build a DOCX version of the research draft from the markdown source."""

from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path("/Users/brendynbeasley/Desktop/mn_climate_project")
DEFAULT_SOURCE = ROOT / "docs" / "mn_climate_economy_research_draft.md"
DEFAULT_OUTPUT = ROOT / "docs" / "mn_climate_economy_research_draft.docx"


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def apply_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def build_docx(source_path: Path, output_path: Path):
    lines = source_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_page_number(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    title_lines = []
    idx = 0
    while idx < len(lines) and len(title_lines) < 4:
        line = lines[idx].strip()
        if line:
            if line.startswith("# "):
                line = line[2:].strip()
            elif line.startswith("## "):
                line = line[3:].strip()
            title_lines.append(line)
        idx += 1

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(title_lines[0])
    apply_run_font(run, size=16, bold=True)
    set_paragraph_spacing(title, after=10, line=1.0)

    for line in title_lines[1:]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        apply_run_font(r, size=12)
        set_paragraph_spacing(p, after=2, line=1.0)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    in_equation = False
    eq_lines = []
    para_buffer = []

    def flush_paragraph():
        nonlocal para_buffer
        if not para_buffer:
            return
        text = " ".join(s.strip() for s in para_buffer).strip()
        if text:
            p = doc.add_paragraph()
            r = p.add_run(text)
            apply_run_font(r, size=12)
            set_paragraph_spacing(p, after=8, line=1.15)
        para_buffer = []

    def flush_equation():
        nonlocal eq_lines
        if not eq_lines:
            return
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(" ".join(line.strip() for line in eq_lines))
        apply_run_font(r, name="Cambria Math", size=11)
        set_paragraph_spacing(p, before=4, after=8, line=1.0)
        eq_lines = []

    while idx < len(lines):
        raw = lines[idx]
        line = raw.rstrip()
        stripped = line.strip()

        if stripped == "$$":
            flush_paragraph()
            if in_equation:
                flush_equation()
                in_equation = False
            else:
                in_equation = True
            idx += 1
            continue

        if in_equation:
            if stripped:
                eq_lines.append(stripped)
            idx += 1
            continue

        if stripped.startswith("# "):
            flush_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(stripped[2:].strip())
            apply_run_font(r, size=14, bold=True)
            set_paragraph_spacing(p, before=8, after=8, line=1.0)
        elif stripped.startswith("## "):
            flush_paragraph()
            p = doc.add_paragraph()
            r = p.add_run(stripped[3:].strip())
            apply_run_font(r, size=13, bold=True)
            set_paragraph_spacing(p, before=10, after=6, line=1.0)
        elif stripped.startswith("### "):
            flush_paragraph()
            p = doc.add_paragraph()
            r = p.add_run(stripped[4:].strip())
            apply_run_font(r, size=12, bold=True)
            set_paragraph_spacing(p, before=8, after=4, line=1.0)
        elif stripped.startswith("- "):
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(stripped[2:].strip())
            apply_run_font(r, size=12)
            set_paragraph_spacing(p, after=2, line=1.1)
        elif not stripped:
            flush_paragraph()
        else:
            para_buffer.append(stripped)
        idx += 1

    flush_paragraph()
    flush_equation()
    doc.save(output_path)
    print(f"Wrote {output_path}")


if __name__ == "__main__":
    source = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_SOURCE
    output = Path(sys.argv[2]) if len(sys.argv) > 2 else DEFAULT_OUTPUT
    build_docx(source, output)
